from pdf2docx import Converter
import docx
import pymorphy3
from docx.shared import RGBColor
from docx2pdf import convert

def change_text_color(doc_file, word, color):
    doc = docx.Document(doc_file)
    for paragraph in doc.paragraphs:
        words = paragraph.text.split()
        for i in range(len(words)):
            if words[i] == word:
                run = paragraph.runs[i]
                font = run.font
                font.color.rgb = RGBColor(*color)

def important_changer(filename, param):
    morph = pymorphy3.MorphAnalyzer()
    cv = Converter(filename)
    cv.convert("Doc1.docx", start=0, end=None)
    cv.close()
    doc = docx.Document('Doc1.docx')
    key_words = ['обязан', 'штраф','расторгнут','одностороннем','передаваться']
    text = []
    att = []
    par = []
    key = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
        words = []
        for i in paragraph.text.split():
            words.append(morph.parse(i)[0][2].lower())
        par.append(words)
    doc = docx.Document()
    doc1 = docx.Document()
    for j in range(len(text)):
        for i in key_words:
            if morph.parse(i.lower())[0][2] in par[j] and text[j] not in att:
                att.append(text[j])
                paragraph = doc.add_paragraph()
                paragraph1 = doc1.add_paragraph()
                sentence_split = text[j].split()
                for x in sentence_split:
                    if x == text[j].split()[par[j].index(morph.parse(i.lower())[0][2])]:
                        run2 = paragraph.add_run(x)
                        run3 = paragraph1.add_run(x)
                        red = RGBColor(255, 0, 0)
                        run2.font.color.rgb = red
                        run3.font.color.rgb = red
                    else:
                        paragraph.add_run(x)
                        paragraph1.add_run(x)
                    paragraph.add_run(" ")
                    paragraph1.add_run(" ")
            else:
                paragraph1 = doc1.add_paragraph()
                paragraph1.add_run(text[j])
                paragraph1.add_run(" ")

    if param != 'highlight':
        doc.save(f'{filename}.docx')
    else:
        doc1.save(f'{filename}.docx')
